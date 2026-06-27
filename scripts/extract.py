#!/usr/bin/env python3
"""
hw-formal-suite/scripts/extract.py

Extract verification requirements from a hardware specification document
(PDF or Word .docx) and output a YAML requirements file.

Usage:
  python scripts/extract.py spec.pdf
  python scripts/extract.py spec.docx --dut apb3_slave --mode formal
  python scripts/extract.py spec.pdf --generate          # chain into generate.py
  python scripts/extract.py docs/*.pdf --output requirements/
"""

import argparse
import base64
import os
import re
import subprocess
import sys
import yaml
import anthropic
from pathlib import Path

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

SUITE_ROOT = Path(__file__).parent.parent
INSTRUCTIONS_PATH = SUITE_ROOT / ".github" / "copilot-instructions.md"

# ── Prompt sent to Claude to extract requirements ────────────────────────────

EXTRACT_SYSTEM = """\
You are an expert hardware verification engineer.
Your task is to read a hardware specification document and extract
all verifiable requirements as a structured YAML file for the
hw-formal-suite SVA generator.

Output format — produce EXACTLY this YAML structure and nothing else:

=== FILE: <stem>_requirements.yaml ===
# Auto-extracted from: <original filename>
name: <dut_name>_check
dut_module: <dut_module_name>
mode: formal          # formal | simulator
package: <pkg>        # apb3 | axi4 | axi3 | apb4 | apb5 | cache | fifo | ... | leave blank if custom

signals:
  clk:   <clock signal name in DUT>
  rst_n: <reset signal name in DUT>
  # list all interface signals as  assertion_port_name: dut_signal_name

parameters:
  # list all configurable parameters with typical values

requirements:
  # One requirement per line, written as a clear, testable statement.
  # Each requirement should map to one or more SVA properties.
  # Be specific: mention signal names, timing (cycles), and conditions.
  - <requirement 1>
  - <requirement 2>
  ...

output:
  dir: generated/<dut_module_name>/
  wrapper: true
=== END FILE ===

Rules:
- Extract ONLY testable functional requirements (timing, protocol, invariants).
- Ignore marketing text, block diagrams, register descriptions not related to timing/protocol.
- If the document covers multiple DUTs, output one YAML block per DUT.
- Use the exact signal names from the specification.
- Requirements must be written in English regardless of the document language.
- If the document is in Japanese, translate requirements to English.
"""

EXTRACT_USER_TEMPLATE = """\
Please extract all verifiable hardware requirements from the attached specification.
{hints}
Output YAML blocks using the === FILE / === END FILE === format.
"""

# ── Document loading ─────────────────────────────────────────────────────────

def load_pdf_as_b64(path: Path) -> str:
    return base64.standard_b64encode(path.read_bytes()).decode("utf-8")


def load_docx_as_text(path: Path) -> str:
    if not HAS_DOCX:
        print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)
    doc = DocxDocument(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Include table content
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


# ── Claude API call ──────────────────────────────────────────────────────────

def extract_from_pdf(
    path: Path,
    client: anthropic.Anthropic,
    model: str,
    hints: str,
) -> str:
    pdf_b64 = load_pdf_as_b64(path)
    user_text = EXTRACT_USER_TEMPLATE.format(hints=hints)

    message = client.messages.create(
        model=model,
        max_tokens=8192,
        system=EXTRACT_SYSTEM,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_b64,
                    },
                },
                {
                    "type": "text",
                    "text": user_text,
                },
            ],
        }],
    )
    return message.content[0].text


def extract_from_docx(
    path: Path,
    client: anthropic.Anthropic,
    model: str,
    hints: str,
) -> str:
    doc_text = load_docx_as_text(path)
    user_text = (
        EXTRACT_USER_TEMPLATE.format(hints=hints)
        + f"\n\n--- DOCUMENT CONTENT ---\n{doc_text}\n--- END ---"
    )

    message = client.messages.create(
        model=model,
        max_tokens=8192,
        system=EXTRACT_SYSTEM,
        messages=[{"role": "user", "content": user_text}],
    )
    return message.content[0].text


# ── Parse output ─────────────────────────────────────────────────────────────

FILE_PATTERN = re.compile(
    r'=== FILE: (.+?) ===\n(.*?)=== END FILE ===',
    re.DOTALL,
)


def parse_yaml_blocks(text: str) -> dict[str, str]:
    blocks = {}
    for m in FILE_PATTERN.finditer(text):
        name    = m.group(1).strip()
        content = m.group(2).strip()
        blocks[name] = content
    return blocks


def validate_yaml(content: str, filename: str) -> bool:
    try:
        data = yaml.safe_load(content)
        required = {"name", "dut_module", "requirements"}
        missing  = required - set(data.keys() if isinstance(data, dict) else [])
        if missing:
            print(f"  Warning: {filename} missing keys: {missing}")
            return False
        return True
    except yaml.YAMLError as e:
        print(f"  Warning: {filename} is not valid YAML: {e}")
        return False


# ── Main ─────────────────────────────────────────────────────────────────────

def process_one(
    spec_path: Path,
    client: anthropic.Anthropic,
    model: str,
    output_dir: Path,
    dut: str,
    mode: str,
    verbose: bool,
) -> list[Path]:
    suffix = spec_path.suffix.lower()

    hints_parts = []
    if dut:
        hints_parts.append(f"The primary DUT module name is: {dut}")
    if mode:
        hints_parts.append(f"Target environment: {mode}")
    hints = "\n".join(hints_parts)

    if verbose:
        print(f"  [{spec_path.name}] Reading document and calling Claude API...")

    if suffix == ".pdf":
        raw = extract_from_pdf(spec_path, client, model, hints)
    elif suffix in {".docx", ".doc"}:
        raw = extract_from_docx(spec_path, client, model, hints)
    else:
        print(f"  ✗  {spec_path.name}: unsupported format (only .pdf and .docx)")
        return []

    blocks = parse_yaml_blocks(raw)

    if not blocks:
        debug = SUITE_ROOT / f".debug_extract_{spec_path.stem}.txt"
        debug.write_text(raw)
        print(f"  △  {spec_path.name}: no YAML blocks found — raw response → {debug}")
        return []

    output_dir.mkdir(parents=True, exist_ok=True)
    saved: list[Path] = []

    for name, content in blocks.items():
        out_path = output_dir / name
        valid    = validate_yaml(content, name)
        out_path.write_text(content + "\n")
        status = "✓" if valid else "△"
        print(f"  {status}  {spec_path.name}  →  {out_path}")
        saved.append(out_path)

    return saved


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Extract SVA requirements from PDF/DOCX spec and output YAML",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python scripts/extract.py docs/apb_spec.pdf
  python scripts/extract.py docs/apb_spec.pdf --dut apb3_slave --mode formal
  python scripts/extract.py docs/apb_spec.pdf --generate        # full pipeline
  python scripts/extract.py docs/*.pdf --output requirements/
        """,
    )
    parser.add_argument("specs", nargs="+", metavar="SPEC",
                        help="PDF or DOCX specification file(s)")
    parser.add_argument("--output", "-o", default="requirements",
                        help="Output directory for YAML files (default: requirements/)")
    parser.add_argument("--dut",
                        help="DUT module name hint (optional)")
    parser.add_argument("--mode", choices=["formal", "simulator"], default="",
                        help="Target environment hint (optional)")
    parser.add_argument("--model", default="claude-sonnet-4-6",
                        help="Claude model ID (default: claude-sonnet-4-6)")
    parser.add_argument("--generate", action="store_true",
                        help="Chain into generate.py after extraction")
    parser.add_argument("--parallel", "-j", type=int, default=1,
                        help="Parallel jobs for --generate step")
    parser.add_argument("--api-key",
                        help="Anthropic API key (or set ANTHROPIC_API_KEY env var)")
    parser.add_argument("--verbose", "-v", action="store_true")
    args = parser.parse_args()

    api_key = args.api_key or os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("Error: ANTHROPIC_API_KEY not set.", file=sys.stderr)
        sys.exit(1)

    client      = anthropic.Anthropic(api_key=api_key)
    output_dir  = Path(args.output)
    spec_paths  = [Path(p) for p in args.specs]

    missing = [p for p in spec_paths if not p.exists()]
    if missing:
        for p in missing:
            print(f"Error: {p} not found", file=sys.stderr)
        sys.exit(1)

    print(f"hw-formal-suite extract — {len(spec_paths)} spec(s), model: {args.model}")
    print(f"Output: {output_dir}/")
    print()

    all_yaml: list[Path] = []
    for spec_path in spec_paths:
        saved = process_one(spec_path, client, args.model, output_dir,
                            args.dut, args.mode, args.verbose)
        all_yaml.extend(saved)

    print()
    print(f"Extracted: {len(all_yaml)} YAML file(s)")

    if args.generate and all_yaml:
        print()
        print("Chaining into generate.py ...")
        print()
        cmd = [
            sys.executable,
            str(SUITE_ROOT / "scripts" / "generate.py"),
            *[str(p) for p in all_yaml],
            "--output", str(SUITE_ROOT / "generated"),
            "--parallel", str(args.parallel),
        ]
        if args.verbose:
            cmd.append("--verbose")
        subprocess.run(cmd, check=False)


if __name__ == "__main__":
    main()
